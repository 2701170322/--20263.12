# -*- coding: utf-8 -*-
"""
生成公租房管理系统软件说明书（软著申请用）Word文档
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ============================================================
# 全局样式设置
# ============================================================
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 页面设置 A4
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.17)
section.right_margin = Cm(3.17)


def set_cell_shading(cell, color):
    """设置单元格底色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading_styled(text, level=1):
    """添加带样式的标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '黑体'
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return heading


def add_para(text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, font_size=12, font_name='宋体'):
    """添加段落"""
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return p


def add_screenshot_placeholder(caption):
    """添加截图占位区域"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'【此处插入{caption}的运行截图】')
    run.font.name = '宋体'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.italic = True
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    # 图片占位框
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('（图片区域）')
    run2.font.name = '宋体'
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(180, 180, 180)
    run2.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    # 图注
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(f'图 {caption}')
    run3.font.name = '宋体'
    run3.font.size = Pt(10)
    run3.bold = True
    run3.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def add_table_row(table, cells_data, header=False):
    """添加表格行"""
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(text))
        run.font.name = '宋体'
        run.font.size = Pt(10)
        run.bold = header
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        if header:
            set_cell_shading(cell, 'D9E2F3')
    return row


# ============================================================
# 封面页
# ============================================================
for _ in range(6):
    doc.add_paragraph()

add_para('公租房管理系统', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font_size=26, font_name='黑体')
doc.add_paragraph()
add_para('软件说明书', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font_size=22, font_name='黑体')

for _ in range(4):
    doc.add_paragraph()

info_lines = [
    '软件名称：公租房管理系统',
    '软件版本：V1.0',
    '编制日期：2026年5月',
]
for line in info_lines:
    add_para(line, align=WD_ALIGN_PARAGRAPH.CENTER, font_size=14)

doc.add_page_break()

# ============================================================
# 目录页（手动目录，因为docx自动目录需要Word刷新）
# ============================================================
add_heading_styled('目  录', level=1)
doc.add_paragraph()

toc_items = [
    ('1', '软件概述'),
    ('1.1', '软件简介'),
    ('1.2', '软件用途'),
    ('1.3', '运行环境'),
    ('2', '软件安装与配置'),
    ('2.1', '安装准备'),
    ('2.2', '数据库配置'),
    ('2.3', '后端服务部署'),
    ('2.4', '前端部署'),
    ('3', '功能说明'),
    ('3.1', '用户登录'),
    ('3.2', '租户注册'),
    ('3.3', '仪表盘'),
    ('3.4', '数据分析'),
    ('3.5', '小区管理'),
    ('3.6', '楼栋管理'),
    ('3.7', '房源管理'),
    ('3.8', '租户管理'),
    ('3.9', '合同管理'),
    ('3.10', '缴费管理'),
    ('3.11', '维修管理'),
    ('3.12', '用户管理'),
    ('3.13', '操作日志'),
    ('3.14', '个人中心'),
    ('3.15', '我的合同'),
    ('3.16', '我的缴费'),
    ('4', '技术特点'),
    ('5', '总结'),
]

for num, title in toc_items:
    p = doc.add_paragraph()
    indent_level = num.count('.')
    p.paragraph_format.left_indent = Cm(indent_level * 1.0)
    run = p.add_run(f'{num}  {title}')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# ============================================================
# 1. 软件概述
# ============================================================
add_heading_styled('1  软件概述', level=1)

add_heading_styled('1.1  软件简介', level=2)
add_para(
    '公租房管理系统是一款面向公共租赁住房管理机构的综合业务管理平台，'
    '采用前后端分离的B/S架构设计。系统基于Vue 3 + Element Plus前端框架与FastAPI后端框架开发，'
    '使用MySQL关系型数据库进行数据持久化存储，通过JWT令牌实现安全可靠的身份认证与权限控制。'
)
add_para(
    '系统涵盖小区管理、楼栋管理、房源管理、租户管理、合同管理、缴费管理、维修管理、用户管理、'
    '操作日志等核心业务模块，同时提供仪表盘和数据分析功能，支持数据可视化展示，'
    '能够全面满足公租房管理机构的日常办公与业务管理需求。'
)

add_heading_styled('1.2  软件用途', level=2)
add_para(
    '本软件主要用于公共租赁住房的日常运营管理，具体用途包括：'
)
purposes = [
    '对公租房小区及楼栋信息进行统一管理，实现房源的精细化管理；',
    '对租户信息进行登记与审核，保障租户资质合规；',
    '管理租赁合同的全生命周期，包括签约、续约、终止等；',
    '实现租金、押金、水电费等费用的收缴与记录；',
    '提供报修工单的全流程管理，包括报修、派单、维修、评价；',
    '通过仪表盘和数据分析，提供运营数据的可视化展示与决策支持；',
    '通过操作日志记录所有用户操作，实现业务审计与追溯；',
    '支持管理员、员工、租户三种角色的权限隔离，确保系统安全。'
]
for p_text in purposes:
    p = doc.add_paragraph(p_text, style='List Bullet')
    for run in p.runs:
        run.font.name = '宋体'
        run.font.size = Pt(12)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

add_heading_styled('1.3  运行环境', level=2)
add_para('1.3.1  硬件环境', bold=True)
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0]
for i, text in enumerate(['项目', '要求']):
    hdr.cells[i].text = ''
    p = hdr.cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(10)
    run.bold = True
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    set_cell_shading(hdr.cells[i], 'D9E2F3')

hw_data = [
    ('处理器', 'Intel Core i5 及以上或同等性能处理器'),
    ('内存', '8GB 及以上'),
    ('硬盘', '100GB 及以上可用空间'),
    ('网络', '支持TCP/IP协议的网络环境'),
]
for row_data in hw_data:
    add_table_row(table, row_data)

doc.add_paragraph()
add_para('1.3.2  软件环境', bold=True)
table2 = doc.add_table(rows=1, cols=2)
table2.style = 'Table Grid'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr2 = table2.rows[0]
for i, text in enumerate(['项目', '要求']):
    hdr2.cells[i].text = ''
    p = hdr2.cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(10)
    run.bold = True
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    set_cell_shading(hdr2.cells[i], 'D9E2F3')

sw_data = [
    ('操作系统', 'Windows 10/11、Linux（CentOS 7+/Ubuntu 18.04+）'),
    ('数据库', 'MySQL 8.0 及以上版本'),
    ('Python', 'Python 3.11 及以上版本'),
    ('Node.js', 'Node.js 18 及以上版本'),
    ('浏览器', 'Chrome 80+、Firefox 80+、Edge 80+'),
    ('Web服务器', 'Uvicorn / Gunicorn（生产环境）'),
]
for row_data in sw_data:
    add_table_row(table2, row_data)

doc.add_page_break()

# ============================================================
# 2. 软件安装与配置
# ============================================================
add_heading_styled('2  软件安装与配置', level=1)

add_heading_styled('2.1  安装准备', level=2)
add_para(
    '在安装公租房管理系统之前，需要确保以下软件已经正确安装并运行：'
)
steps = [
    'MySQL 8.0 数据库服务已安装并正常运行；',
    'Python 3.11 及以上版本已安装，并配置好环境变量；',
    'Node.js 18 及以上版本已安装，并配置好环境变量；',
    '操作系统已安装必要的网络组件，支持TCP/IP通信。'
]
for i, step in enumerate(steps, 1):
    add_para(f'({i}) {step}')

add_heading_styled('2.2  数据库配置', level=2)
add_para(
    '使用MySQL客户端工具连接数据库服务器，创建系统所需的数据库，执行以下SQL语句：'
)
p = doc.add_paragraph()
run = p.add_run(
    'CREATE DATABASE public_housing_db CHARACTER SET utf8mb4 COLLATE utf8mb4_unicode_ci;'
)
run.font.name = 'Consolas'
run.font.size = Pt(10)

add_para(
    '然后编辑后端配置文件 backend/.env，填写数据库连接信息，包括数据库主机地址、端口、用户名、密码及数据库名。'
)

add_heading_styled('2.3  后端服务部署', level=2)
add_para(
    '进入后端目录 backend，安装Python依赖包，执行数据库初始化脚本，然后启动后端服务：'
)
commands = [
    'pip install -r requirements.txt   # 安装Python依赖',
    'python init_db.py                  # 初始化数据库及默认管理员账号',
    'python run.py                      # 启动后端服务（默认端口8001）',
]
for cmd in commands:
    p = doc.add_paragraph()
    run = p.add_run(cmd)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

add_para(
    '后端服务启动后，可通过访问 http://localhost:8001/docs 查看API接口文档，'
    '访问 http://localhost:8001/health 进行健康检查。'
)

add_heading_styled('2.4  前端部署', level=2)
add_para(
    '进入前端目录 frontend，安装Node.js依赖包并启动前端开发服务器：'
)
fe_commands = [
    'npm install      # 安装前端依赖',
    'npm run dev      # 启动开发服务器（默认端口3000）',
    'npm run build    # 打包为生产环境静态文件',
]
for cmd in fe_commands:
    p = doc.add_paragraph()
    run = p.add_run(cmd)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

add_para(
    '开发模式下，前端通过Vite开发服务器的代理配置将API请求转发到后端8001端口。'
    '生产环境下，将打包生成的dist目录部署到Nginx等Web服务器即可。'
)

doc.add_page_break()

# ============================================================
# 3. 功能说明
# ============================================================
add_heading_styled('3  功能说明', level=1)
add_para(
    '公租房管理系统采用基于角色的访问控制机制，提供管理员、员工、租户三种角色，'
    '不同角色拥有不同的功能权限。以下对各功能模块进行详细说明。'
)

# --- 3.1 用户登录 ---
add_heading_styled('3.1  用户登录', level=2)
add_para(
    '用户登录是系统的入口功能。用户在登录页面输入用户名和密码，系统验证用户身份信息，'
    '验证通过后根据用户角色跳转到对应的功能首页：管理员和员工进入仪表盘，租户进入个人中心。'
    '系统采用JWT（JSON Web Token）令牌机制进行身份认证，登录成功后服务端返回访问令牌，'
    '前端在后续请求中携带该令牌进行身份验证，令牌有效期为30分钟。'
)
add_para('功能要点：', bold=True)
login_features = [
    '支持用户名/密码方式登录；',
    '登录成功后根据角色自动跳转对应首页；',
    '采用JWT令牌进行安全认证；',
    '登录失败时给出明确的错误提示信息；',
    '支持"记住我"功能，延长登录状态有效期。'
]
for f in login_features:
    p = doc.add_paragraph(f, style='List Bullet')
    for run in p.runs:
        run.font.name = '宋体'
        run.font.size = Pt(12)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

add_screenshot_placeholder('用户登录页面')

# --- 3.2 租户注册 ---
add_heading_styled('3.2  租户注册', level=2)
add_para(
    '租户注册功能允许新的租户用户自助注册系统账号。注册时需填写用户名、密码、确认密码、'
    '真实姓名、身份证号、手机号等必要信息。注册提交后，账号状态默认为"待审核"，'
    '需要管理员在租户管理模块中审核通过后方可正常登录使用系统。'
    '该机制有效防止了非法用户进入系统，保障了公租房管理的安全性。'
)
add_screenshot_placeholder('租户注册页面')

# --- 3.3 仪表盘 ---
add_heading_styled('3.3  仪表盘', level=2)
add_para(
    '仪表盘是管理员和员工登录后的默认首页，以数据可视化方式展示系统的核心运营指标，'
    '帮助管理人员快速掌握公租房运营状况。仪表盘包含以下数据展示区域：'
)
dashboard_items = [
    '数据概览卡片：显示小区总数、房源总数、租户总数、合同总数等核心统计指标；',
    '房源状态统计：以图表形式展示空闲、已租、维修中、预留等各状态房源的数量与占比；',
    '缴费统计：展示已缴、未缴、逾期等缴费状态分布情况；',
    '维修工单统计：展示待处理、进行中、已完成等各状态维修工单数量；',
    '近期动态：展示最近的操作记录和业务变动情况。'
]
for item in dashboard_items:
    p = doc.add_paragraph(item, style='List Bullet')
    for run in p.runs:
        run.font.name = '宋体'
        run.font.size = Pt(12)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

add_screenshot_placeholder('仪表盘页面')

# --- 3.4 数据分析 ---
add_heading_styled('3.4  数据分析', level=2)
add_para(
    '数据分析模块提供更深入的数据统计与可视化展示，支持多维度数据分析。'
    '该模块面向管理人员，提供以下分析功能：'
)
analysis_items = [
    '房源分布分析：按小区、楼栋维度展示房源分布情况；',
    '租金收入趋势：按时间维度展示租金收入的变化趋势；',
    '租户构成分析：按收入水平、家庭规模等维度分析租户构成；',
    '合同到期预警：展示即将到期的合同列表，便于及时续约处理；',
    '缴费率分析：统计各小区、各时段的缴费率情况。'
]
for item in analysis_items:
    p = doc.add_paragraph(item, style='List Bullet')
    for run in p.runs:
        run.font.name = '宋体'
        run.font.size = Pt(12)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

add_screenshot_placeholder('数据分析页面')

# --- 3.5 小区管理 ---
add_heading_styled('3.5  小区管理', level=2)
add_para(
    '小区管理模块用于管理公租房小区的基本信息，包括小区名称、地址、楼栋总数、单元总数、'
    '负责人姓名及联系电话等。系统支持小区信息的新增、修改、删除和查询操作，'
    '并提供分页列表展示和关键词搜索功能。'
)
add_para('功能要点：', bold=True)
community_features = [
    '新增小区：填写小区名称、地址、负责人等基本信息；',
    '编辑小区：修改已有小区的各项信息；',
    '删除小区：删除不再管理的小区（关联楼栋时提示确认）；',
    '查询搜索：支持按小区名称等关键词进行模糊搜索；',
    '列表展示：分页展示小区信息，显示关键字段。'
]
for f in community_features:
    p = doc.add_paragraph(f, style='List Bullet')
    for run in p.runs:
        run.font.name = '宋体'
        run.font.size = Pt(12)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

add_screenshot_placeholder('小区管理-列表页面')
add_screenshot_placeholder('小区管理-新增/编辑对话框')

# --- 3.6 楼栋管理 ---
add_heading_styled('3.6  楼栋管理', level=2)
add_para(
    '楼栋管理模块用于管理各小区下的楼栋信息，包括楼栋编号、所属小区、楼层数、单元数等。'
    '楼栋与小区为从属关系，每个小区可包含多栋楼。系统支持楼栋信息的增删改查操作，'
    '新增楼栋时需选择所属小区，确保数据的层级关系正确。'
)
add_screenshot_placeholder('楼栋管理-列表页面')
add_screenshot_placeholder('楼栋管理-新增/编辑对话框')

# --- 3.7 房源管理 ---
add_heading_styled('3.7  房源管理', level=2)
add_para(
    '房源管理是系统的核心模块之一，用于管理公租房的具体房源信息。'
    '每套房源关联到所属楼栋，包含单元号、楼层、面积、房间数、月租金、押金、状态、'
    '配套设施及描述等信息。'
)
add_para('房源状态说明：', bold=True)
status_table = doc.add_table(rows=1, cols=3)
status_table.style = 'Table Grid'
status_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = status_table.rows[0]
for i, text in enumerate(['状态', '标识', '说明']):
    hdr.cells[i].text = ''
    p = hdr.cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(10)
    run.bold = True
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    set_cell_shading(hdr.cells[i], 'D9E2F3')

status_data = [
    ('空闲', 'vacant', '房源当前无人入住，可分配给租户'),
    ('已租', 'occupied', '房源已被租户承租，处于合同期内'),
    ('维修中', 'maintenance', '房源正在维修，暂不可出租'),
    ('预留', 'reserved', '房源已被预留，等待特定租户签约'),
]
for row_data in status_data:
    add_table_row(status_table, row_data)

doc.add_paragraph()
add_para(
    '房源状态的变更与合同签订、终止及维修工单等业务操作联动，'
    '确保房源状态始终与实际业务情况保持一致。'
)
add_screenshot_placeholder('房源管理-列表页面')
add_screenshot_placeholder('房源管理-新增/编辑对话框')

# --- 3.8 租户管理 ---
add_heading_styled('3.8  租户管理', level=2)
add_para(
    '租户管理模块用于管理公租房租户的详细信息，包括姓名、身份证号、手机号、邮箱、'
    '住址、紧急联系人、职业、收入水平、家庭人口数、审核状态等。'
)
add_para('租户状态说明：', bold=True)
tenant_status = [
    ('待审核', 'pending', '租户已注册但尚未通过管理员审核'),
    ('已激活', 'active', '租户审核通过，可正常登录使用系统'),
    ('已停用', 'inactive', '租户账号被停用，无法登录系统'),
]
ts_table = doc.add_table(rows=1, cols=3)
ts_table.style = 'Table Grid'
ts_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = ts_table.rows[0]
for i, text in enumerate(['状态', '标识', '说明']):
    hdr.cells[i].text = ''
    p = hdr.cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(10)
    run.bold = True
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    set_cell_shading(hdr.cells[i], 'D9E2F3')
for row_data in tenant_status:
    add_table_row(ts_table, row_data)

doc.add_paragraph()
add_para(
    '管理员可在租户管理模块中对租户进行审核操作，审核通过后租户方可正常登录系统。'
    '同时支持对租户信息的编辑、删除和查询操作。'
)
add_screenshot_placeholder('租户管理-列表页面')
add_screenshot_placeholder('租户管理-审核操作')

# --- 3.9 合同管理 ---
add_heading_styled('3.9  合同管理', level=2)
add_para(
    '合同管理模块用于管理公租房租赁合同的全生命周期。合同信息包括合同编号、'
    '关联租户、关联房源、起止日期、月租金、押金金额、缴费日、是否自动续约、状态等。'
)
add_para('合同状态说明：', bold=True)
contract_status = [
    ('待生效', 'pending', '合同已创建但尚未到起始日期'),
    ('生效中', 'active', '合同处于有效期内，正在执行'),
    ('已到期', 'expired', '合同已超过结束日期，自然到期'),
    ('已终止', 'terminated', '合同在到期前被主动终止'),
]
cs_table = doc.add_table(rows=1, cols=3)
cs_table.style = 'Table Grid'
cs_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = cs_table.rows[0]
for i, text in enumerate(['状态', '标识', '说明']):
    hdr.cells[i].text = ''
    p = hdr.cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(10)
    run.bold = True
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    set_cell_shading(hdr.cells[i], 'D9E2F3')
for row_data in contract_status:
    add_table_row(cs_table, row_data)

doc.add_paragraph()
add_para(
    '合同签订时，系统自动将关联房源状态更新为"已租"；'
    '合同终止时，系统自动将房源状态恢复为"空闲"，确保房源状态与合同状态同步。'
)
add_screenshot_placeholder('合同管理-列表页面')
add_screenshot_placeholder('合同管理-新增/编辑对话框')
add_screenshot_placeholder('合同管理-终止合同操作')

# --- 3.10 缴费管理 ---
add_heading_styled('3.10  缴费管理', level=2)
add_para(
    '缴费管理模块用于管理租户的各类费用收缴情况，包括租金、押金、水电费、维修费等。'
    '每条缴费记录关联到具体租户和合同，记录缴费类型、应缴金额、已缴金额、到期日期、'
    '缴费日期、缴费方式、交易流水号等信息。'
)
add_para('缴费类型说明：', bold=True)
pay_types = [
    ('租金', 'rent', '每月应缴纳的房屋租金'),
    ('押金', 'deposit', '入住时缴纳的押金'),
    ('水电费', 'utilities', '每月产生的水电费用'),
    ('维修费', 'maintenance', '维修产生的费用'),
    ('其他', 'other', '其他类型的费用'),
]
pt_table = doc.add_table(rows=1, cols=3)
pt_table.style = 'Table Grid'
pt_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = pt_table.rows[0]
for i, text in enumerate(['类型', '标识', '说明']):
    hdr.cells[i].text = ''
    p = hdr.cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(10)
    run.bold = True
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    set_cell_shading(hdr.cells[i], 'D9E2F3')
for row_data in pay_types:
    add_table_row(pt_table, row_data)

doc.add_paragraph()
add_para('缴费状态说明：', bold=True)
pay_status = [
    ('未缴费', 'unpaid', '费用尚未缴纳'),
    ('已缴费', 'paid', '费用已全额缴纳'),
    ('部分缴费', 'partial', '费用已部分缴纳'),
    ('逾期', 'overdue', '费用已超过缴费期限未缴纳'),
]
ps_table = doc.add_table(rows=1, cols=3)
ps_table.style = 'Table Grid'
ps_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = ps_table.rows[0]
for i, text in enumerate(['状态', '标识', '说明']):
    hdr.cells[i].text = ''
    p = hdr.cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(10)
    run.bold = True
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    set_cell_shading(hdr.cells[i], 'D9E2F3')
for row_data in pay_status:
    add_table_row(ps_table, row_data)

doc.add_paragraph()
add_para(
    '管理员或员工可通过缴费管理模块进行收款操作，记录实际缴费金额、缴费方式和交易流水号。'
)
add_screenshot_placeholder('缴费管理-列表页面')
add_screenshot_placeholder('缴费管理-收款操作')

# --- 3.11 维修管理 ---
add_heading_styled('3.11  维修管理', level=2)
add_para(
    '维修管理模块提供报修工单的全流程管理功能，涵盖报修提交、工单指派、维修执行、完工确认和评价反馈。'
    '每条维修记录关联到具体房源和租户，记录工单标题、描述、优先级、状态、预计费用、实际费用、'
    '图片附件、评价反馈和评分等信息。'
)
add_para('维修工单状态流转：', bold=True)
maint_status = [
    ('待处理', 'pending', '工单已提交，等待指派维修人员'),
    ('已指派', 'assigned', '工单已指派给维修人员，等待开工'),
    ('维修中', 'in_progress', '维修人员正在进行维修作业'),
    ('已完成', 'completed', '维修作业完成，租户可评价'),
    ('已取消', 'cancelled', '工单被取消，不再处理'),
]
ms_table = doc.add_table(rows=1, cols=3)
ms_table.style = 'Table Grid'
ms_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = ms_table.rows[0]
for i, text in enumerate(['状态', '标识', '说明']):
    hdr.cells[i].text = ''
    p = hdr.cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(10)
    run.bold = True
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    set_cell_shading(hdr.cells[i], 'D9E2F3')
for row_data in maint_status:
    add_table_row(ms_table, row_data)

doc.add_paragraph()
add_para('优先级说明：', bold=True)
maint_priority = [
    ('低', 'low', '不影响正常居住的轻微问题'),
    ('中', 'medium', '有一定影响但不紧急的问题'),
    ('高', 'high', '影响正常居住的重要问题'),
    ('紧急', 'urgent', '严重影响安全或居住的紧急问题'),
]
mp_table = doc.add_table(rows=1, cols=3)
mp_table.style = 'Table Grid'
mp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = mp_table.rows[0]
for i, text in enumerate(['优先级', '标识', '说明']):
    hdr.cells[i].text = ''
    p = hdr.cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(10)
    run.bold = True
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    set_cell_shading(hdr.cells[i], 'D9E2F3')
for row_data in maint_priority:
    add_table_row(mp_table, row_data)

doc.add_paragraph()
add_screenshot_placeholder('维修管理-工单列表页面')
add_screenshot_placeholder('维修管理-新增报修工单')
add_screenshot_placeholder('维修管理-指派维修人员')
add_screenshot_placeholder('维修管理-完工确认与评价')

# --- 3.12 用户管理 ---
add_heading_styled('3.12  用户管理', level=2)
add_para(
    '用户管理模块仅管理员可访问，用于管理系统中的用户账号。支持新增用户、编辑用户信息、'
    '删除用户、启用/停用用户等操作。新增用户时需指定用户名、密码和角色（管理员/员工/租户）。'
    '系统提供用户列表展示，支持按用户名和角色进行筛选查询。'
)
add_para('角色权限说明：', bold=True)
role_table = doc.add_table(rows=1, cols=3)
role_table.style = 'Table Grid'
role_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = role_table.rows[0]
for i, text in enumerate(['角色', '标识', '权限范围']):
    hdr.cells[i].text = ''
    p = hdr.cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(10)
    run.bold = True
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    set_cell_shading(hdr.cells[i], 'D9E2F3')

role_data = [
    ('管理员', 'admin', '全部功能模块，包括用户管理和操作日志'),
    ('员工', 'staff', '业务管理模块，不含用户管理和操作日志'),
    ('租户', 'tenant', '个人中心、我的合同、我的缴费'),
]
for row_data in role_data:
    add_table_row(role_table, row_data)

doc.add_paragraph()
add_screenshot_placeholder('用户管理-列表页面')
add_screenshot_placeholder('用户管理-新增/编辑对话框')

# --- 3.13 操作日志 ---
add_heading_styled('3.13  操作日志', level=2)
add_para(
    '操作日志模块仅管理员可访问，系统通过操作日志中间件自动记录所有API请求的操作日志，'
    '包括操作用户、操作动作（创建/查询/更新/删除/登录）、操作模块、操作描述、IP地址、'
    '用户代理、请求方法、请求路径、响应状态码、执行时间等详细信息。'
    '管理员可按时间范围、操作模块、操作动作等条件进行筛选查询，实现完整的业务审计与操作追溯。'
)
add_screenshot_placeholder('操作日志-列表页面')

# --- 3.14 个人中心 ---
add_heading_styled('3.14  个人中心', level=2)
add_para(
    '个人中心是租户角色的专属功能页面，租户登录后可在此查看和编辑个人基本信息，'
    '包括姓名、身份证号、手机号、邮箱、住址、紧急联系人等。'
    '个人中心为租户提供了便捷的信息维护入口，确保个人信息的准确性和时效性。'
)
add_screenshot_placeholder('个人中心页面')

# --- 3.15 我的合同 ---
add_heading_styled('3.15  我的合同', level=2)
add_para(
    '我的合同是租户角色的专属功能页面，租户可在此查看自己名下的所有租赁合同信息，'
    '包括合同编号、关联房源、起止日期、月租金、合同状态等。'
    '租户可快速了解当前合同的执行状态和到期时间，便于及时办理续约或退租手续。'
)
add_screenshot_placeholder('我的合同页面')

# --- 3.16 我的缴费 ---
add_heading_styled('3.16  我的缴费', level=2)
add_para(
    '我的缴费是租户角色的专属功能页面，租户可在此查看自己的所有缴费记录，'
    '包括缴费类型、应缴金额、已缴金额、到期日期、缴费状态等。'
    '租户可快速了解自己的缴费情况，及时处理未缴费和逾期费用。'
)
add_screenshot_placeholder('我的缴费页面')

doc.add_page_break()

# ============================================================
# 4. 技术特点
# ============================================================
add_heading_styled('4  技术特点', level=1)

tech_features = [
    ('前后端分离架构', 
     '系统采用前后端分离的B/S架构，前端使用Vue 3 + Element Plus构建响应式用户界面，'
     '后端使用FastAPI提供RESTful API服务，前后端通过HTTP协议通信，'
     '实现了界面展示与业务逻辑的完全解耦，便于独立开发、测试和部署。'),
    ('JWT安全认证', 
     '系统采用JWT（JSON Web Token）令牌机制进行身份认证，用户登录成功后服务端签发访问令牌，'
     '前端在后续请求的Authorization头部携带该令牌，后端中间件自动验证令牌有效性，'
     '实现了无状态的安全认证机制，有效防止了CSRF等安全威胁。'),
    ('基于角色的访问控制', 
     '系统实现了管理员、员工、租户三种角色的细粒度权限控制，'
     '前端根据用户角色动态渲染菜单和路由，后端接口层面进行角色权限校验，'
     '双重保障确保不同角色只能访问其权限范围内的功能和数据。'),
    ('全链路操作审计', 
     '系统通过操作日志中间件自动记录所有API请求的详细信息，包括操作用户、动作类型、'
     '操作模块、请求参数、响应状态、执行时间等，实现了完整的操作审计与追溯能力，'
     '满足公租房管理的合规性要求。'),
    ('数据可视化分析', 
     '系统提供仪表盘和数据分析两大可视化模块，通过图表组件直观展示房源状态分布、'
     '缴费统计、维修工单统计、租金收入趋势等关键运营指标，'
     '帮助管理人员快速掌握运营状况，辅助决策分析。'),
    ('业务状态联动', 
     '系统实现了房源状态、合同状态、缴费状态等核心业务状态之间的自动联动，'
     '如合同签订时自动更新房源为"已租"，合同终止时自动恢复房源为"空闲"，'
     '确保了数据的一致性和业务的正确性。'),
]

for title, desc in tech_features:
    add_para(title, bold=True)
    add_para(desc)

doc.add_page_break()

# ============================================================
# 5. 总结
# ============================================================
add_heading_styled('5  总结', level=1)
add_para(
    '公租房管理系统是一款功能完善、架构合理、安全可靠的公共租赁住房综合管理平台。'
    '系统采用当前主流的前后端分离架构，基于Vue 3和FastAPI框架开发，'
    '具有界面友好、操作便捷、性能优越、扩展性强等特点。'
)
add_para(
    '系统涵盖小区管理、楼栋管理、房源管理、租户管理、合同管理、缴费管理、维修管理、'
    '用户管理、操作日志等核心业务模块，同时提供数据可视化和分析功能，'
    '能够全面满足公租房管理机构的日常办公与业务管理需求。'
)
add_para(
    '在安全方面，系统实现了JWT身份认证、基于角色的访问控制、全链路操作审计等安全机制，'
    '确保了系统数据的安全性和业务操作的合规性。'
    '在技术方面，系统采用了前后端分离、RESTful API设计、ORM数据访问等现代化开发模式，'
    '保证了系统的可维护性和可扩展性。'
)
add_para(
    '综上所述，公租房管理系统具备完善的业务功能、良好的用户体验和可靠的安全保障，'
    '能够为公租房管理机构提供高效、便捷的信息化管理工具，具有重要的实用价值和推广意义。'
)

# ============================================================
# 保存文档
# ============================================================
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '公租房管理系统-软件说明书.docx')
doc.save(output_path)
print(f'软件说明书已生成: {output_path}')
