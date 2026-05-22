# -*- coding: utf-8 -*-
"""
生成公租房管理系统源代码文档（软著申请用）
规则：源代码不足60页则全部提交，超过60页则提交前30页+后30页（每页50行）
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ============================================================
# 配置：源代码文件列表（按逻辑顺序排列）
# ============================================================
BASE_DIR = r"C:\Users\王康\public_housing_management"

SOURCE_FILES = [
    # 后端 - 入口与配置
    "backend/run.py",
    "backend/init_db.py",
    "backend/app/config.py",
    "backend/app/database.py",
    "backend/app/auth.py",
    "backend/app/models.py",
    "backend/app/schemas.py",
    "backend/app/main.py",
    # 后端 - 路由模块
    "backend/app/routers/auth.py",
    "backend/app/routers/users.py",
    "backend/app/routers/communities.py",
    "backend/app/routers/buildings.py",
    "backend/app/routers/houses.py",
    "backend/app/routers/tenants.py",
    "backend/app/routers/contracts.py",
    "backend/app/routers/payments.py",
    "backend/app/routers/dashboard.py",
    "backend/app/routers/maintenance.py",
    "backend/app/routers/logs.py",
    # 前端 - 入口与配置
    "frontend/src/main.js",
    "frontend/src/App.vue",
    "frontend/src/api/index.js",
    "frontend/src/router/index.js",
    "frontend/src/stores/auth.js",
    "frontend/src/layouts/MainLayout.vue",
    # 前端 - 视图页面
    "frontend/src/views/Login.vue",
    "frontend/src/views/Register.vue",
    "frontend/src/views/Dashboard.vue",
    "frontend/src/views/DataAnalysis.vue",
    "frontend/src/views/Communities.vue",
    "frontend/src/views/Buildings.vue",
    "frontend/src/views/Houses.vue",
    "frontend/src/views/Tenants.vue",
    "frontend/src/views/Contracts.vue",
    "frontend/src/views/Payments.vue",
    "frontend/src/views/Maintenance.vue",
    "frontend/src/views/Users.vue",
    "frontend/src/views/Logs.vue",
    "frontend/src/views/Profile.vue",
    "frontend/src/views/MyContracts.vue",
    "frontend/src/views/MyPayments.vue",
    "frontend/src/views/NotFound.vue",
]

LINES_PER_PAGE = 50
MAX_PAGES = 60
HALF_PAGES = 30


def read_source_code():
    """读取所有源代码文件，返回 (文件路径, 代码行列表) 的列表"""
    all_code = []
    for rel_path in SOURCE_FILES:
        abs_path = os.path.join(BASE_DIR, rel_path)
        if not os.path.exists(abs_path):
            print(f"  [跳过] 文件不存在: {rel_path}")
            continue
        try:
            with open(abs_path, 'r', encoding='utf-8') as f:
                lines = f.readlines()
        except UnicodeDecodeError:
            with open(abs_path, 'r', encoding='gbk') as f:
                lines = f.readlines()

        # 添加文件分隔标记
        all_code.append(("__FILE_HEADER__", rel_path))
        for i, line in enumerate(lines, 1):
            # 保留换行符
            all_code.append((rel_path, line.rstrip('\n\r')))
        # 文件末尾空行
        all_code.append(("__FILE_END__", ""))

    return all_code


def select_code_lines(all_code):
    """
    根据软著要求选取代码行
    总行数 <= 3000行(60页): 全部提交
    总行数 > 3000行: 前30页 + 后30页
    """
    # 只计算实际代码行（排除文件头尾标记）
    code_lines = [(path, line) for path, line in all_code if path not in ("__FILE_HEADER__", "__FILE_END__")]
    total_lines = len(code_lines)
    total_pages = (total_lines + LINES_PER_PAGE - 1) // LINES_PER_PAGE

    print(f"  源代码总行数: {total_lines}")
    print(f"  折合页数: {total_pages}")

    if total_pages <= MAX_PAGES:
        print(f"  不足{MAX_PAGES}页，提交全部代码")
        return all_code

    # 超过60页：取前30页 + 后30页
    first_half = LINES_PER_PAGE * HALF_PAGES  # 前1500行
    last_half = LINES_PER_PAGE * HALF_PAGES   # 后1500行

    first_lines = code_lines[:first_half]
    last_lines = code_lines[-last_half:]

    print(f"  提取前{HALF_PAGES}页({first_half}行) + 后{HALF_PAGES}页({last_half}行)")
    print(f"  省略中间 {total_lines - first_half - last_half} 行")

    return first_lines, last_lines


def generate_docx(output_path):
    """生成源代码Word文档"""
    doc = Document()

    # ============================================================
    # 全局样式
    # ============================================================
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(10)
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 页面设置 A4
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ============================================================
    # 封面页
    # ============================================================
    for _ in range(8):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('公租房管理系统')
    run.font.name = '黑体'
    run.font.size = Pt(26)
    run.bold = True
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('源代码文档')
    run.font.name = '黑体'
    run.font.size = Pt(22)
    run.bold = True
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    for _ in range(6):
        doc.add_paragraph()

    info_lines = [
        '软件名称：公租房管理系统',
        '软件版本：V1.0',
        '编制日期：2026年5月',
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.name = '宋体'
        run.font.size = Pt(14)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_page_break()

    # ============================================================
    # 读取源代码
    # ============================================================
    print("正在读取源代码文件...")
    all_code = read_source_code()
    result = select_code_lines(all_code)

    # 判断是全量还是前后分段
    if isinstance(result, tuple):
        first_lines, last_lines = result
        need_split = True
    else:
        first_lines = result
        last_lines = []
        need_split = False

    # ============================================================
    # 写入代码
    # ============================================================
    page_num = 1
    line_num = 0
    current_file = None

    def add_code_line(file_path, line_text, line_number, page, is_omitted=False):
        nonlocal current_file
        # 文件切换时添加分隔注释
        if file_path != current_file:
            if current_file is not None:
                # 空行分隔
                sep_p = doc.add_paragraph()
                sep_p.paragraph_format.line_spacing = 1.0
                sep_p.paragraph_format.space_before = Pt(0)
                sep_p.paragraph_format.space_after = Pt(0)
                run = sep_p.add_run('')
                run.font.size = Pt(10)

            # 文件路径注释
            comment_p = doc.add_paragraph()
            comment_p.paragraph_format.line_spacing = 1.0
            comment_p.paragraph_format.space_before = Pt(0)
            comment_p.paragraph_format.space_after = Pt(0)

            # 左侧文件路径
            run1 = comment_p.add_run(f'# {file_path}')
            run1.font.name = 'Consolas'
            run1.font.size = Pt(9)
            run1.font.color.rgb = RGBColor(0, 100, 180)
            run1.bold = True
            run1.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            current_file = file_path

        # 代码行
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

        # 行号
        line_str = f'{line_number:>4}  '
        run_num = p.add_run(line_str)
        run_num.font.name = 'Consolas'
        run_num.font.size = Pt(9)
        run_num.font.color.rgb = RGBColor(150, 150, 150)
        run_num.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 代码内容
        run_code = p.add_run(line_text if line_text else ' ')
        run_code.font.name = 'Consolas'
        run_code.font.size = Pt(9)
        run_code.font.color.rgb = RGBColor(30, 30, 30)
        run_code.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    def write_code_block(code_lines, start_page=1):
        nonlocal page_num, line_num, current_file
        page_num = start_page
        line_num = 0
        current_file = None

        for file_path, line_text in code_lines:
            line_num += 1

            # 每50行换页
            if line_num > LINES_PER_PAGE and (line_num - 1) % LINES_PER_PAGE == 0:
                doc.add_page_break()
                page_num += 1

            add_code_line(file_path, line_text, line_num, page_num)

    # 写入前半部分
    print("正在生成Word文档...")
    write_code_block(first_lines, start_page=1)

    # 如果需要分段，插入省略提示后写入后半部分
    if need_split:
        # 省略提示
        doc.add_page_break()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for _ in range(8):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('( 以下为源代码最后30页 )')
        run.font.name = '黑体'
        run.font.size = Pt(14)
        run.bold = True
        run.font.color.rgb = RGBColor(128, 128, 128)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        doc.add_page_break()

        # 计算后半部分的起始行号
        total_code = len(first_lines) + len(last_lines)
        # 实际总行数需要加回被省略的中间部分
        # 获取实际总行数
        code_all = [(path, line) for path, line in all_code if path not in ("__FILE_HEADER__", "__FILE_END__")]
        actual_total = len(code_all)
        start_line = actual_total - len(last_lines) + 1

        line_num = 0
        current_file = None
        page_num = HALF_PAGES + 1  # 续接页码

        # 重新计算行号（从 actual_total - last_half + 1 开始）
        for file_path, line_text in last_lines:
            line_num += 1
            if line_num > LINES_PER_PAGE and (line_num - 1) % LINES_PER_PAGE == 0:
                doc.add_page_break()
                page_num += 1

            add_code_line(file_path, line_text, start_line + line_num - 1, page_num)

    # ============================================================
    # 页脚（页码）
    # ============================================================
    # 为每个section添加页码
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加页码字段
        run = p.add_run()
        run.font.name = '宋体'
        run.font.size = Pt(9)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 使用Word域代码插入页码
        fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fld_char_begin)

        run2 = p.add_run()
        run2.font.name = '宋体'
        run2.font.size = Pt(9)
        instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._r.append(instr)

        run3 = p.add_run()
        run3.font.name = '宋体'
        run3.font.size = Pt(9)
        fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._r.append(fld_char_end)

    # 保存
    doc.save(output_path)
    print(f"\n源代码文档已生成: {output_path}")
    print(f"  总页数约: {page_num}")


if __name__ == '__main__':
    output = os.path.join(BASE_DIR, '公租房管理系统-源代码文档.docx')
    generate_docx(output)
